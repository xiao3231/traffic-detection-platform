#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将 docs/毕业设计论文-正文.md 转为 Word。运行：python docs/build_thesis_word.py"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
MD = ROOT / "毕业设计论文-正文.md"
OUT = ROOT / "青岛理工大学-论文正文（系统生成）.docx"


def font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    text = re.sub(r"<sup>(.*?)</sup>", r"\1", text)
    # 处理上标引用 [1][2]
    parts = re.split(r"(\[\d+\])", text)
    for part in parts:
        if re.fullmatch(r"\[\d+\]", part):
            r = p.add_run(part)
            font(r, size=12)
            r.font.superscript = True
        else:
            r = p.add_run(part)
            font(r, size=12)


def heading(doc, text, level):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, "黑体", {1: 16, 2: 14, 3: 12}.get(level, 12), True)


def table_from_md(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=cols)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.rows[i].cells[j].text = cell.strip()


def main():
    text = MD.read_text(encoding="utf-8")
    doc = Document()
    lines = text.splitlines()
    i = 0
    table_buf = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        if line.startswith("# 毕业设计") or line.startswith("> "):
            i += 1
            continue
        if line.startswith("---"):
            i += 1
            continue
        if line.startswith("|") and "|" in line[1:]:
            if re.match(r"^\|[-:\s|]+\|$", line):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_buf = []
            table_buf.append([c.strip() for c in line.strip("|").split("|")])
            i += 1
            continue
        elif in_table:
            table_from_md(doc, table_buf)
            in_table = False
            table_buf = []

        if line.startswith("## "):
            heading(doc, line[3:].strip(), 1)
        elif line.startswith("### "):
            heading(doc, line[4:].strip(), 2)
        elif line.startswith("**图") or line.startswith("【图") or line.startswith("【请"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip("*"))
            font(r, "宋体", 10.5)
        elif line.startswith("**表"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip("*"))
            font(r, "宋体", 10.5)
        elif line.strip():
            body(doc, line.strip())
        i += 1

    if table_buf:
        table_from_md(doc, table_buf)

    doc.save(OUT)
    print("已生成:", OUT)


if __name__ == "__main__":
    main()
