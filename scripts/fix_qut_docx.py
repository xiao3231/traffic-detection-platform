#!/usr/bin/env python3
"""Fix 青岛理工大学.docx: clean header, task-book layout, table row integrity."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

DOC_PATH = Path("/Users/neptune/Downloads/青岛理工大学.docx")
HEADER_TEXT = "青岛理工大学毕业设计（论文）"
PAGE_MARKERS = {"第1页", "第2页", "第3页"}


def set_bottom_border(paragraph, sz: str = "12", color: str = "000000") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    for old in p_pr.findall(qn("w:pBdr")):
        p_pr.remove(old)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def rebuild_header(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.header_distance = Pt(12)
    section.top_margin = Pt(72)

    for header in (section.header, section.first_page_header):
        for paragraph in list(header.paragraphs):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    paragraph = section.header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(HEADER_TEXT)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    set_bottom_border(paragraph)

    # Cover / first page: no header content
    first = section.first_page_header
    if first.paragraphs:
        clear_paragraph(first.paragraphs[0])
    else:
        first.add_paragraph()


def set_row_cant_split(table) -> None:
    for row in table.rows:
        tr = row._tr
        tr_pr = tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            cant = OxmlElement("w:cantSplit")
            cant.set(qn("w:val"), "true")
            tr_pr.append(cant)


def remove_page_marker_paragraphs(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in PAGE_MARKERS:
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1
    return removed


def insert_page_break_before_table(doc: Document, table_index: int) -> None:
    table = doc.tables[table_index]
    tbl_element = table._element
    prev = tbl_element.getprevious()
    if prev is not None and prev.tag == qn("w:p"):
        texts = [node.text or "" for node in prev.iter(qn("w:t"))]
        if any("PAGE" in (t or "") for t in texts):
            pass
        # If previous paragraph already ends with page break, skip
        if prev.findall(".//" + qn("w:br")):
            br_types = [
                br.get(qn("w:type"))
                for br in prev.findall(".//" + qn("w:br"))
            ]
            if "page" in br_types:
                return

    body = doc.element.body
    idx = list(body).index(tbl_element)
    if idx == 0:
        return
    prev_element = body[idx - 1]
    if prev_element.tag == qn("w:p"):
        paragraph = doc.paragraphs[[p._element for p in doc.paragraphs].index(prev_element)]
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    else:
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r.append(br)
        p.append(r)
        tbl_element.addprevious(p)


def format_task_book_heading(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in ("毕业设计(论文)任务书", "毕业设计（论文）任务书"):
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("毕业设计（论文）任务书")
            run.bold = True
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(16)
            break

    for paragraph in doc.paragraphs:
        if "专业班级" in paragraph.text and "学号" in paragraph.text:
            clear_paragraph(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(
                "专业班级  网安223    学号  202201050847    "
                "姓名  肖萌    下发日期  2026.5"
            )
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(12)
            break


def main() -> None:
    backup = DOC_PATH.with_suffix(".docx.bak")
    shutil.copy2(DOC_PATH, backup)
    print(f"Backup: {backup}")

    doc = Document(str(DOC_PATH))
    rebuild_header(doc)
    removed = remove_page_marker_paragraphs(doc)
    print(f"Removed page markers: {removed}")

    for table in doc.tables:
        set_row_cant_split(table)

    # Keep 评语 / 答辩 tables on separate pages without manual “第N页” text
    for idx in (1, 2, 3):
        if idx < len(doc.tables):
            insert_page_break_before_table(doc, idx)

    format_task_book_heading(doc)
    doc.save(str(DOC_PATH))
    print(f"Saved: {DOC_PATH}")


if __name__ == "__main__":
    main()
